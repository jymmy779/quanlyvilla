"""Generate a .docx report from the ProjectReport.md markdown file.

The script reads the markdown content, performs a very lightweight conversion
to Word using the ``python-docx`` library, and writes the output to
``ProjectReport.docx`` in the repository root.

The conversion rules are simple and aim to preserve the structure of the
original document:
* ``#``       → Heading level 1 (Title)
* ``##``      → Heading level 2 (Section)
* ``###``     → Heading level 3 (Sub‑section)
* ``- ``      → Bullet list item
* ``|``       → Table rows – the script attempts to create a table for the
                first occurrence of a markdown table (used for the team list)
                and then treats subsequent lines as plain paragraphs.
* All other lines are added as normal paragraphs.

This implementation is sufficient for the provided ``ProjectReport.md`` and
produces a readable Word document without requiring a full markdown parser.
"""

import pathlib
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

ROOT = pathlib.Path(__file__).parent
MD_PATH = ROOT / "ProjectReport.md"
# Output a uniquely named file to avoid conflicts with existing documents
DOCX_PATH = ROOT / "ProjectReport_filled.docx"


def is_table_separator(line: str) -> bool:
    """Return True if the line is a markdown table separator like ``|---|``.
    """
    stripped = line.strip()
    return stripped.startswith("|") and set(stripped) <= {"|", "-", ":", " "}


def add_heading(paragraph, level: int):
    """Apply heading style based on level (1‑3)."""
    if level == 1:
        paragraph.style = "Heading 1"
    elif level == 2:
        paragraph.style = "Heading 2"
    elif level == 3:
        paragraph.style = "Heading 3"
    else:
        paragraph.style = "Normal"


def main():
    if not MD_PATH.exists():
        raise FileNotFoundError(f"Markdown source not found: {MD_PATH}")

    doc = Document()
    # Use a default font size for body text
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    with MD_PATH.open(encoding="utf-8") as f:
        lines = [ln.rstrip("\n") for ln in f]

    i = 0
    while i < len(lines):
        line = lines[i]
        # Headings
        if line.startswith("# "):
            p = doc.add_paragraph(line[2:].strip())
            add_heading(p, 1)
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph(line[3:].strip())
            add_heading(p, 2)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:].strip())
            add_heading(p, 3)
            i += 1
            continue
        # Horizontal rule
        if line.strip().startswith("---"):
            doc.add_paragraph().add_run().add_break()
            i += 1
            continue
        # Bullet list
        if line.lstrip().startswith("- "):
            p = doc.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
            i += 1
            continue
        # Table – only process the first markdown table (team list)
        if line.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            # Gather header row and subsequent rows until a blank line or non‑table line
            header = [h.strip() for h in line.strip("|").split("|")]
            rows = []
            i += 2  # skip separator line
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip("|").split("|")])
                i += 1
            # Create table in docx
            table = doc.add_table(rows=1, cols=len(header))
            hdr_cells = table.rows[0].cells
            for idx, txt in enumerate(header):
                hdr_cells[idx].text = txt
            for row in rows:
                row_cells = table.add_row().cells
                for idx, txt in enumerate(row):
                    if idx < len(row_cells):
                        row_cells[idx].text = txt
            continue
        # Empty line -> add paragraph break
        if line.strip() == "":
            doc.add_paragraph()
            i += 1
            continue
        # Default: normal paragraph
        doc.add_paragraph(line)
        i += 1

    doc.save(DOCX_PATH)
    print(f"Report generated: {DOCX_PATH}")


if __name__ == "__main__":
    main()
