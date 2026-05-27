import docx
import sys

def main():
    docx_path = r"C:\Users\thaiv\Downloads\Mau Cuon Bao Cao Do An CNPM.docx"
    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        print(f"Error opening docx: {e}", file=sys.stderr)
        return
        
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print(f"Number of tables: {len(doc.tables)}")
    
    with open("template_text.txt", "w", encoding="utf-8") as f:
        f.write("=== PARAGRAPHS ===\n")
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            style = p.style.name
            if text:
                f.write(f"[{style}] P{i}: {text}\n")
                
        f.write("\n=== TABLES ===\n")
        for t_idx, t in enumerate(doc.tables):
            f.write(f"\n--- TABLE {t_idx} ---\n")
            for r_idx, r in enumerate(t.rows):
                row_cells = [c.text.strip() for c in r.cells]
                f.write(f"Row {r_idx}: {' | '.join(row_cells)}\n")

if __name__ == '__main__':
    main()
